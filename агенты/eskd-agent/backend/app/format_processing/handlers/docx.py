"""Извлечение текста из DOCX."""

from __future__ import annotations

import io

from docx import Document


def extract_docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        line = para.text.strip()
        if line:
            parts.append(line)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()

"""Извлечение текста из вложений (PDF, DOCX, XLSX, TXT, изображения OCR)."""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import PurePath

from agent_pochta.schemas import Attachment

SUPPORTED_MIME = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "text/plain",
    "image/jpeg",
    "image/png",
    "image/gif",
    "application/zip",
    "application/x-zip-compressed",
}

_EXT_TO_MIME = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".txt": "text/plain",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".gif": "image/gif",
    ".zip": "application/zip",
}


def resolve_mime_type(attachment: Attachment) -> str:
    """MIME из заголовка письма или по расширению файла."""
    mime = (attachment.mime_type or "").split(";", 1)[0].strip().lower()
    if mime and mime not in ("application/octet-stream", "binary/octet-stream"):
        return mime
    ext = PurePath(attachment.filename or "").suffix.lower()
    return _EXT_TO_MIME.get(ext, mime or "application/octet-stream")


def is_supported_attachment(attachment: Attachment) -> bool:
    return resolve_mime_type(attachment) in SUPPORTED_MIME


def normalize_extracted_text(text: str, *, max_chars: int) -> str:
    """Сжимает пробелы и обрезает слишком длинный текст для LLM/RAG."""
    cleaned = re.sub(r"[ \t]+\n", "\n", text)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    cleaned = cleaned.strip()
    if len(cleaned) <= max_chars:
        return cleaned
    truncated = cleaned[: max_chars - 1]
    if " " in truncated[-80:]:
        truncated = truncated.rsplit(" ", 1)[0]
    return f"{truncated}…"


def is_meaningful_extracted_text(text: str | None) -> bool:
    """True, если извлечён реальный текст (не служебная заглушка)."""
    if not text or not text.strip():
        return False
    stripped = text.strip()
    if stripped.startswith("[") and (
        "не извлечён" in stripped
        or "OCR не распознал" in stripped
        or "Заглушка извлечения" in stripped
    ):
        return False
    return True


def extract_pdf(content: bytes) -> str:
    import fitz  # pymupdf

    doc = fitz.open(stream=content, filetype="pdf")
    try:
        parts: list[str] = []
        for page in doc:
            page_text = page.get_text().strip()
            if page_text:
                parts.append(page_text)
        return "\n\n".join(parts)
    finally:
        doc.close()


def extract_docx(content: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(content))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_xlsx(content: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    try:
        parts: list[str] = []
        for sheet in workbook.worksheets:
            sheet_lines: list[str] = []
            for row in sheet.iter_rows(values_only=True):
                cells = [
                    str(value).strip()
                    for value in row
                    if value is not None and str(value).strip()
                ]
                if cells:
                    sheet_lines.append(" | ".join(cells))
            if sheet_lines:
                parts.append(f"[Лист {sheet.title}]\n" + "\n".join(sheet_lines))
        return "\n\n".join(parts)
    finally:
        workbook.close()


def extract_plain_text(content: bytes) -> str:
    for encoding in ("utf-8", "cp1251", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def extract_image_ocr(content: bytes) -> str:
    from PIL import Image
    import pytesseract

    image = Image.open(BytesIO(content))
    try:
        return pytesseract.image_to_string(image, lang="rus+eng")
    except pytesseract.TesseractNotFoundError as exc:
        raise RuntimeError("Tesseract OCR не установлен в системе") from exc
    except pytesseract.TesseractError:
        return pytesseract.image_to_string(image)


def extract_zip(content: bytes, *, max_chars: int) -> str:
    """Извлекает текст из первых поддерживаемых файлов внутри ZIP-архива."""
    import zipfile

    parts: list[str] = []
    budget = max_chars
    with zipfile.ZipFile(BytesIO(content)) as archive:
        for info in archive.infolist():
            if info.is_dir() or budget <= 0:
                continue
            name = PurePath(info.filename).name
            if not name or name.startswith("."):
                continue
            ext = PurePath(name).suffix.lower()
            if ext == ".zip":
                continue
            inner_mime = _EXT_TO_MIME.get(ext, "application/octet-stream")
            if inner_mime not in SUPPORTED_MIME or inner_mime in {
                "application/zip",
                "application/x-zip-compressed",
            }:
                continue
            try:
                inner_bytes = archive.read(info)
            except Exception:
                continue
            inner_att = Attachment(
                filename=name,
                mime_type=inner_mime,
                size_bytes=len(inner_bytes),
                content=inner_bytes,
            )
            text, _ocr = extract_attachment_text(inner_att, max_chars=budget)
            if text:
                parts.append(f"[Файл в архиве {info.filename}]\n{text}")
                budget = max(0, budget - len(text))
    return "\n\n".join(parts)


def extract_attachment_text(
    attachment: Attachment,
    *,
    max_chars: int = 12_000,
) -> tuple[str | None, bool]:
    """Извлекает текст из вложения. Возвращает (text, ocr_used)."""
    if attachment.content is None:
        return None, False

    mime = resolve_mime_type(attachment)
    if mime not in SUPPORTED_MIME:
        return None, False

    raw = ""
    ocr_used = False
    try:
        if mime == "application/pdf":
            raw = extract_pdf(attachment.content)
        elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            raw = extract_docx(attachment.content)
        elif mime == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            raw = extract_xlsx(attachment.content)
        elif mime == "text/plain":
            raw = extract_plain_text(attachment.content)
        elif mime in {"application/zip", "application/x-zip-compressed"}:
            raw = extract_zip(attachment.content, max_chars=max_chars)
        elif mime.startswith("image/"):
            try:
                raw = extract_image_ocr(attachment.content)
                ocr_used = True
            except (ImportError, OSError, RuntimeError):
                return None, False
    except ImportError as exc:
        raise ImportError(
            'Установите парсер вложений: pip install -e ".[documents]"'
        ) from exc
    except Exception:
        return None, ocr_used

    raw = raw.strip()
    if not raw:
        return None, ocr_used
    return normalize_extracted_text(raw, max_chars=max_chars), ocr_used

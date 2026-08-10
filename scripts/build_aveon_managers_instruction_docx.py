"""Generate Word instruction for Avion procurement agent (managers & supervisors)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path(__file__).resolve().parents[1] / "data" / "aveon"
DOC_PATH = OUT_DIR / "Инструкция_Авион_менеджеры_и_руководители.docx"


def add_bold_runs(paragraph, text: str) -> None:
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:
            run.bold = True


def build_docx() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    styles["Normal"].font.size = Pt(11)

    title = doc.add_heading("Агент закупок «Авион»", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Инструкция для менеджеров и руководителя")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    doc.add_heading("Как открыть", level=1)
    for item in [
        "Войти в платформу AI Agent.",
        "В каталоге агентов выбрать «Агент закупок (Авион)» (на экране — «Закупки · Авион»).",
        "После первого анализа данные сохраняются: при следующем входе дашборды и задания подгрузятся автоматически.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("Менеджер по закупкам", level=1)
    doc.add_heading("Ежедневная работа", level=2)
    for item in [
        "Откройте агента «Закупки · Авион».",
        "Перейдите в блок «Мои задания» (или вкладку «Задания» на дашборде обеспеченности).",
        "Сначала обработайте задания в разделе «Срочные» (приоритеты «Срочно» и «Сегодня»), затем — «На неделю».",
        "Типы заданий: Отгрузка; Логистика МСК; Таможня; Логистика Ростов; Необходимые закупки.",
        "Выполните действие по заданию и заполните поле «Результат» (результат работы).",
        "После заполнения система проверит результат и покажет статус: выполнено / частично / не выполнено.",
        "При необходимости используйте подсказку ИИ для формулировки результата.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("Завершение смены", level=2)
    for item in [
        "В конце рабочего дня нажмите «Завершить смену».",
        "Укажите причины по незакрытым заданиям.",
        "Подтвердите отправку — отчёт уйдёт руководителю на e-mail.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("Что видит менеджер", level=2)
    for item in [
        "Только свои задания (по колонке «Ответственный менеджер»).",
        "Дашборд обеспеченности: изделия, номенклатуры, задания.",
        "Доску рисков логистики (контрольные точки: загрузка → МСК → таможня → Ростов).",
        "Блок «Результаты работы менеджеров» менеджер не видит.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Руководитель", level=1)
    doc.add_heading("Контроль после анализа", level=2)
    for item in [
        "Откройте агента «Закупки · Авион» (после того как аналитик провёл анализ Excel).",
        "Проверьте дашборд обеспеченности — вкладки «Изделия», «Номенклатуры», «Задания»; период: день / неделя / месяц.",
        "Просмотрите доску рисков логистики — фильтр по уровню риска.",
        "Откройте «Сменное задание» — полный список заданий по всем менеджерам.",
        "Скачайте файл «сменное_задание_закупки.xlsx» при необходимости.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("Контроль работы менеджеров", level=2)
    for item in [
        "На дашборде обеспеченности откройте блок «Результаты работы менеджеров».",
        "Смотрите по каждому менеджеру: всего заданий / выполнено / не закрыто.",
        "Читайте причины незакрытых задач из отчётов о завершении смены.",
        "Отчёты о завершении смены также приходят на e-mail (тема: «Avion: отчёт завершения смены …»).",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("Что видит руководитель дополнительно", level=2)
    for item in [
        "Полное «Сменное задание» по всем менеджерам (не «Мои задания»).",
        "Сводку «Результаты работы менеджеров».",
        "Те же дашборды обеспеченности и логистики, что и у менеджеров.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Краткий чеклист", level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Роль"
    headers[1].text = "Действие"
    for cell in headers:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    table.rows[1].cells[0].text = "Менеджер"
    table.rows[1].cells[1].text = (
        "Открыть агент → «Мои задания» → выполнить → заполнить «Результат» → «Завершить смену»"
    )
    table.rows[2].cells[0].text = "Руководитель"
    table.rows[2].cells[1].text = (
        "Открыть агент → обеспеченность и риски → «Сменное задание» → «Результаты работы менеджеров»"
    )

    doc.add_heading("Обратная связь", level=1)
    doc.add_paragraph(
        "В правом нижнем углу — виджет «Обратная связь по агенту закупок Авион». "
        "Используйте его для замечаний и предложений по работе агента."
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOC_PATH)


def main() -> None:
    build_docx()
    print(DOC_PATH)


if __name__ == "__main__":
    main()

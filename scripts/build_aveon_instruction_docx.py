"""Generate Word instruction + flowchart for Aveon procurement agent."""

from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUT_DIR = Path(__file__).resolve().parents[1] / "data" / "aveon"
DOC_PATH = OUT_DIR / "Инструкция_агент_закупок_Авион.docx"
DIAGRAM_PATH = OUT_DIR / "схема_агент_Авион.png"


def try_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    if bold:
        candidates = [
            r"C:\Windows\Fonts\segoeuib.ttf",
            r"C:\Windows\Fonts\arialbd.ttf",
            r"C:\Windows\Fonts\segoeui.ttf",
            r"C:\Windows\Fonts\arial.ttf",
        ]
    else:
        candidates = [
            r"C:\Windows\Fonts\segoeui.ttf",
            r"C:\Windows\Fonts\arial.ttf",
        ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: float) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        test = f"{current} {word}".strip()
        if draw.textlength(test, font=font) <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [text]


def draw_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    text: str,
    fill: str,
    outline: str,
    font,
    padding: int = 10,
    radius: int = 12,
) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=2)
    max_w = x2 - x1 - padding * 2
    lines: list[str] = []
    for part in text.split("\n"):
        lines.extend(wrap_text(draw, part, font, max_w))
    line_h = getattr(font, "size", 16) + 4
    total_h = len(lines) * line_h
    ty = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        tw = draw.textlength(line, font=font)
        tx = x1 + (x2 - x1 - tw) / 2
        draw.text((tx, ty), line, fill="#1a2332", font=font)
        ty += line_h


def draw_arrow(draw: ImageDraw.ImageDraw, x: int, y1: int, y2: int, color: str = "#334155") -> None:
    draw.line((x, y1, x, y2 - 8), fill=color, width=2)
    draw.polygon([(x - 6, y2 - 10), (x + 6, y2 - 10), (x, y2)], fill=color)


def build_diagram() -> None:
    width, height = 1100, 1680
    img = Image.new("RGB", (width, height), "#F7F9FC")
    draw = ImageDraw.Draw(img)
    title_font = try_font(28, bold=True)
    box_font = try_font(17)
    small_font = try_font(15)
    legend_font = try_font(14)
    note_font = try_font(14)

    draw.text((40, 24), "Агент закупок Авион — схема работы", fill="#0F172A", font=title_font)
    draw.text(
        (40, 64),
        "От загрузки Excel до выгрузки result.xlsx",
        fill="#475569",
        font=small_font,
    )

    legend_y = 100
    legends = [
        (40, "#DBEAFE", "#2563EB", "Что грузит пользователь"),
        (300, "#DCFCE7", "#16A34A", "Файлы в data/aveon"),
        (560, "#FEF3C7", "#D97706", "Обработка на backend"),
        (820, "#F3E8FF", "#7C3AED", "Результат / UI"),
    ]
    for x, fill, outline, label in legends:
        draw.rounded_rectangle((x, legend_y, x + 240, legend_y + 36), radius=8, fill=fill, outline=outline, width=2)
        draw.text((x + 14, legend_y + 8), label, fill="#1a2332", font=legend_font)

    cx = width // 2
    box_w = 540
    x1 = cx - box_w // 2
    x2 = cx + box_w // 2

    steps = [
        ("user", "1. Открытие страницы агента\n«Закупки · Авион»"),
        ("user", "2. Загрузка Excel пользователем\n• график производства\n• остатки\n• график отгрузок"),
        ("proc", "3. Нажатие «Анализировать»\nPOST analyze-excel"),
        ("proc", "4. Определение ролей файлов\n(LM Studio / локально)"),
        ("proc", "5. Извлечение изделий и\nпомесячного плана из графика производства"),
        ("root", "6. Сопоставление с data/aveon\n• Сопоставление номенклатур.xlsx\n• Сокол Спецификация из 1с.xlsx"),
        ("proc", "7. Сбор материалов спецификаций\n→ уникальные номенклатуры"),
        ("root", "8. Подстановка цен и поставщиков\nиз «Цены закупки…xlsx»"),
        ("user", "9. Остатки из файла пользователя\n+ потребность по месяцам\nплан изделия × qty из спеки"),
        ("user", "10. Ожидаемые поступления\nи риски логистики\nиз графика отгрузок"),
        ("root", "11. Сборка result.xlsx\nпо шаблону Header.xlsx"),
        ("result", "12. Скачивание итоговой таблицы\n+ доска рисков логистики в UI"),
    ]
    colors = {
        "user": ("#DBEAFE", "#2563EB"),
        "root": ("#DCFCE7", "#16A34A"),
        "proc": ("#FEF3C7", "#D97706"),
        "result": ("#F3E8FF", "#7C3AED"),
    }

    y = 160
    gap = 26
    for i, (kind, text) in enumerate(steps):
        fill, outline = colors[kind]
        lines_count = text.count("\n") + 1
        h = max(78, 24 + lines_count * 22)
        draw_box(draw, (x1, y, x2, y + h), text, fill, outline, box_font)
        if i < len(steps) - 1:
            draw_arrow(draw, cx, y + h, y + h + gap)
        y += h + gap

    draw.rounded_rectangle((40, 200, 250, 340), radius=10, fill="#FFFFFF", outline="#CBD5E1", width=1)
    draw.text((55, 215), "Связь данных", fill="#0F172A", font=try_font(15, bold=True))
    for i, line in enumerate(
        [
            "изделие графика",
            "→ номенклатура",
            "→ лист спецификации",
            "→ материалы",
            "→ цены / остатки",
            "→ потребность",
        ]
    ):
        draw.text((55, 245 + i * 14), line, fill="#334155", font=note_font)

    draw.rounded_rectangle((850, 200, 1060, 320), radius=10, fill="#FFFFFF", outline="#CBD5E1", width=1)
    draw.text((865, 215), "Формула", fill="#0F172A", font=try_font(15, bold=True))
    for i, line in enumerate(
        [
            "demand[месяц] =",
            "Σ план_изделия[месяц]",
            "× qty из спеки",
        ]
    ):
        draw.text((865, 245 + i * 18), line, fill="#334155", font=note_font)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    img.save(DIAGRAM_PATH, "PNG")


def build_docx() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    styles["Normal"].font.size = Pt(11)

    title = doc.add_heading("Агент закупок «Авион»", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        "Краткая инструкция: от открытия страницы до выгрузки итоговой таблицы"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    doc.add_heading("Суть", level=1)
    doc.add_paragraph(
        "Агент считает потребность в материалах для плана производства: "
        "пользователь загружает оперативные Excel (график производства, остатки, график отгрузок), "
        "backend сопоставляет их со справочниками в папке data/aveon и отдаёт итоговый result.xlsx "
        "плюс риски логистики на сегодня."
    )

    doc.add_heading("Что где лежит", level=1)
    doc.add_heading("Загружает пользователь (каждый запуск)", level=2)
    for name, desc in [
        ("График производства", "изделия и план выпуска по месяцам."),
        ("Остатки", "сколько номенклатуры уже есть на складе."),
        (
            "График отгрузок",
            "ожидаемые поступления по датам + данные для рисков логистики "
            "(сроки до МСК / МСК→Ростов).",
        ),
    ]:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(name).bold = True
        para.add_run(f" — {desc}")

    note = doc.add_paragraph()
    note.add_run("Роли файлов агент определяет сам").bold = True
    note.add_run(" (по содержимому), вручную назначать не нужно.")

    doc.add_heading("Лежит в data/aveon (справочники, не грузятся с UI)", level=2)
    for name, desc in [
        ("Сопоставление номенклатур.xlsx", "связь: изделие из графика → номенклатура."),
        (
            "Сокол Спецификация из 1с.xlsx",
            "листы спецификаций: из чего состоит изделие (материалы и qty).",
        ),
        ("Цены закупки за 2026_….xlsx", "поставщик и цена по номенклатуре."),
        ("Header.xlsx", "шаблон колонок итогового result.xlsx (данные с 5-й строки)."),
    ]:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(name).bold = True
        para.add_run(f" — {desc}")

    doc.add_heading("Как связано", level=1)
    doc.add_paragraph(
        "Изделие (график производства) → номенклатура (сопоставление) → лист спецификации → "
        "материалы → цена/поставщик + остаток → потребность и поступления по месяцам → result.xlsx."
    )
    formula = doc.add_paragraph()
    formula.add_run("Потребность: ").bold = True
    formula.add_run("demand[месяц] = Σ (план_изделия[месяц] × qty_материала_из_спеки).")

    doc.add_heading("Пошагово: работа агента", level=1)
    for i, (step_title, body) in enumerate(
        [
            ("Открыть страницу агента", "В интерфейсе платформы — агент «Закупки · Авион»."),
            (
                "Загрузить Excel",
                "Перетащить или выбрать файлы: график производства, остатки, график отгрузок.",
            ),
            ("Нажать «Анализировать»", "Файлы уходят на backend (analyze-excel)."),
            (
                "Определение ролей",
                "Система помечает каждый файл: production_schedule / stock / shipment_schedule (и др.).",
            ),
            ("Изделия из графика", "Читается таблица изделий и помесячные qty."),
            (
                "Сопоставление со спеками",
                "Через файлы в data/aveon находится лист спецификации для каждого изделия.",
            ),
            (
                "Сбор материалов",
                "Из спек собираются уникальные номенклатуры с привязкой к изделиям.",
            ),
            ("Цены и поставщики", "Подставляются из справочника цен в data/aveon."),
            (
                "Остатки и потребность",
                "Остатки — из файла пользователя; потребность считается по месяцам (формула выше).",
            ),
            (
                "Поступления и логистика",
                "Из графика отгрузок — ожидаемые поступления по месяцам; контрольные точки "
                "логистики на сегодня (загрузка → МСК → таможня → Ростов).",
            ),
            ("Формирование result.xlsx", "Таблица собирается по шаблону Header.xlsx."),
            (
                "Выгрузка результата",
                "Браузер скачивает итоговый Excel; на экране — сводка и доска рисков логистики.",
            ),
        ],
        1,
    ):
        para = doc.add_paragraph()
        para.add_run(f"{i}. {step_title}. ").bold = True
        para.add_run(body)

    doc.add_heading("Блок-схема", level=1)
    doc.add_paragraph("Ниже — схема потока от открытия страницы до скачивания итога.")
    doc.add_picture(str(DIAGRAM_PATH), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph("Рисунок. Поток работы агента закупок Авион")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_heading("Короткий чеклист", level=1)
    for item in [
        "В data/aveon лежат актуальные: сопоставление, спецификации, цены, Header.",
        "Пользователь загрузил минимум график производства (желательно ещё остатки и отгрузки).",
        "После анализа скачан result.xlsx; при необходимости проверены риски логистики в UI.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.save(DOC_PATH)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_diagram()
    build_docx()
    print(f"diagram: {DIAGRAM_PATH}")
    print(f"doc: {DOC_PATH}")
    print(f"size: {DOC_PATH.stat().st_size}")


if __name__ == "__main__":
    main()

from __future__ import annotations

import io
from datetime import date

from docx import Document as DocxDocument

from app.models.enums import NdChangeOperationType


class DocxEditor:
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def build_draft(
        self,
        *,
        title: str,
        request_number: str,
        reason: str,
        old_text: str | None,
        new_text: str,
        source_bytes: bytes | None = None,
        operation_type: NdChangeOperationType = NdChangeOperationType.MANUAL_REVIEW,
    ) -> bytes:
        if source_bytes:
            try:
                document = DocxDocument(io.BytesIO(source_bytes))
                replaced = self._replace_in_document(document, old_text, new_text)
                if not replaced:
                    document.add_page_break()
                    document.add_heading("Проект изменения", level=1)
                    document.add_paragraph(f"Заявка: {request_number}")
                    document.add_paragraph(f"Причина: {reason}")
                    document.add_heading("Новая редакция", level=2)
                    document.add_paragraph(new_text)
                return self._to_bytes(document)
            except Exception:
                pass

        document = DocxDocument()
        document.add_heading(f"Проект новой редакции: {title}", level=1)
        document.add_paragraph(f"Заявка: {request_number}")
        document.add_paragraph(f"Причина изменения: {reason}")
        document.add_paragraph(f"Тип операции: {operation_type.value}")
        document.add_heading("Текущая редакция", level=2)
        document.add_paragraph(old_text or "Редактируемый исходник отсутствует или место изменения требует ручной проверки.")
        document.add_heading("Новая редакция", level=2)
        document.add_paragraph(new_text)
        return self._to_bytes(document)

    def build_notice(
        self,
        *,
        request_number: str,
        document_title: str,
        reason: str,
        release_date: date | None,
        effective_date: date | None,
        change_text: str,
        distribution_list: list[str],
        attachments: list[str],
        initiator_comment: str | None,
    ) -> bytes:
        document = DocxDocument()
        document.add_heading(f"Извещение об изменении {request_number}", level=1)
        table = document.add_table(rows=0, cols=2)
        rows = [
            ("Документ", document_title),
            ("Дата выпуска", release_date.isoformat() if release_date else "-"),
            ("Дата введения изменения", effective_date.isoformat() if effective_date else "-"),
            ("Причина изменения", reason),
            ("Разослать", ", ".join(distribution_list) if distribution_list else "-"),
            ("Приложения", ", ".join(attachments) if attachments else "-"),
            ("Комментарий инициатора", initiator_comment or "-"),
        ]
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
        document.add_heading("Содержание изменения", level=2)
        document.add_paragraph(change_text)
        return self._to_bytes(document)

    def _replace_in_document(self, document: DocxDocument, old_text: str | None, new_text: str) -> bool:
        if not old_text:
            return False
        normalized_old = " ".join(old_text.split())
        for paragraph in document.paragraphs:
            if normalized_old and normalized_old in " ".join(paragraph.text.split()):
                paragraph.text = new_text
                return True
        return False

    def _to_bytes(self, document: DocxDocument) -> bytes:
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()
